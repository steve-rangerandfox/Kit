"""
Build a clean Kit feature catalog as a Word .docx.

Run: python scripts/build-features-doc.py
Output: Kit-Features.docx in the repo root.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "Kit-Features.docx"


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    return p


def feature(doc, name, summary, trigger=None, requires=None, notes=None):
    add_h3(doc, name)
    add_para(doc, summary)
    if trigger:
        add_kv(doc, "How to use", trigger)
    if requires:
        add_kv(doc, "Requires", requires)
    if notes:
        add_kv(doc, "Notes", notes)


def divider(doc):
    p = doc.add_paragraph("⸻")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()

    # Doc-wide style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title ───────────────────────────────────────────────
    title = doc.add_heading("Kit — Feature Catalog", level=0)
    sub = doc.add_paragraph()
    sub_r = sub.add_run(f"Ranger & Fox · {date.today().isoformat()}")
    sub_r.italic = True
    sub_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    intro = doc.add_paragraph()
    intro.add_run(
        "Kit is an AI-powered production agent that lives in Slack. It provisions "
        "new projects across multiple services, tracks time, manages files, surfaces "
        "review feedback, transcodes deliverables, captures notes, builds a self-evolving "
        "brain per project channel, flags mistakes and deadlines, and answers questions "
        "about the studio's full history. Everything below is live unless explicitly "
        "marked otherwise."
    )

    # ── 1. Provisioning & Setup ────────────────────────────
    add_h1(doc, "1. Provisioning & Setup")

    feature(
        doc,
        "New project provisioning",
        "Spin up a new project across Slack, Dropbox, Harvest, Frame.io, and Boords "
        "in parallel from one Slack form. Creates the project channel, folder trees, "
        "Harvest project + tasks, Frame.io project + folders, and a starter storyboard. "
        "Streams DM progress while it runs and posts a summary card to the project channel.",
        trigger="/kit newproject in Slack — pick services, fill in details, submit.",
        notes="On completion the new channel auto-receives a Brain canvas (see § 6).",
    )

    feature(
        doc,
        "Storyboards (Boords)",
        "Turns a script (pasted text or .docx/.txt drop) into a Boords storyboard with "
        "shot panels, voiceover, and frame placeholders.",
        trigger="/storyboard in Slack, or DM Kit \"new storyboard\" and drop the script file.",
    )

    feature(
        doc,
        "Freelancer onboarding",
        "One form provisions a freelancer across Slack (channel invite), Dropbox "
        "(folder share), Frame.io (project membership), and Harvest (user creation + "
        "project assignment).",
        trigger="/kit onboard in Slack. Restricted to PMs / CDs / admins.",
    )

    # ── 2. Daily Communication ─────────────────────────────
    add_h1(doc, "2. Daily Communication")

    feature(
        doc,
        "Conversational assistant",
        "@mention Kit in any channel or DM it directly to ask about projects, budgets, "
        "schedules, files, reviews, contacts, or anything in the studio's knowledge base. "
        "Routes queries to the right specialist agent under the hood (Harvest, Dropbox, "
        "Frame.io, Slack, Studio Knowledge, Brain).",
        trigger="@Kit <question> in a channel, or DM Kit directly.",
    )

    feature(
        doc,
        "Daily hours check-in",
        "At 5pm local time, Kit DMs each active team member with a check-in card. Reply "
        "with hours per project and Kit posts them to Harvest. A single nudge fires at "
        "10pm if no reply came in.",
        requires="HARVEST_ACCESS_TOKEN, staff directory populated. Timezone via CHECKIN_TIMEZONE.",
    )

    feature(
        doc,
        "Ad-hoc hours logging",
        "DM Kit with something like \"3 hours on STUDIO100 today, color review\" and Kit "
        "parses it into a Harvest entry. Shows a confirmation card before posting.",
    )

    feature(
        doc,
        "Notes capture",
        "Save freeform notes against a project from anywhere in Slack. Notes get "
        "embedded into the RAG store immediately so Kit can quote from them, and "
        "are also fed into the project's Brain.",
        trigger=(
            "@Kit note for <project>: <body> · "
            "@Kit remember that <body> · "
            "/kit note <project> | <body>"
        ),
    )

    feature(
        doc,
        "Shot list canvas",
        "Build a Boords-style shot list directly inside a Slack channel as a Canvas. "
        "Paste a script and Kit produces structured shot rows; edit by talking to it; "
        "attach reference thumbnails by dropping images in-thread.",
        trigger=(
            "@Kit shot list from this: <script> · "
            "@Kit add a close-up between 2 and 3 · "
            "/kit shotlist <script>"
        ),
        requires="Slack scope canvases:write.",
    )

    # ── 3. Detection & Integration ─────────────────────────
    add_h1(doc, "3. Detection & Integration")

    feature(
        doc,
        "Frame.io review-link detection",
        "When anyone pastes a Frame.io URL in a project channel, Kit posts a card with "
        "asset thumbnails, current comment count, and an approve / request-changes "
        "summary pulled from Frame.io v4 API.",
    )

    feature(
        doc,
        "Dropbox → Frame.io file watcher",
        "New files in a project's Dropbox folder auto-upload to the matching Frame.io "
        "project for review. Runs continuously via cursor-based delta polling.",
    )

    feature(
        doc,
        "Pre-meeting briefings",
        "Reads each watched Google Calendar 30 minutes before a meeting and posts a "
        "briefing to the relevant project channel (and optionally DMs the producer): "
        "attendees, project context, recent activity, open decisions from the Brain.",
        trigger="Automatic 30 min before each calendar event.",
        requires=(
            "GOOGLE_SERVICE_ACCOUNT_JSON, GOOGLE_CALENDAR_IDS, "
            "GOOGLE_CALENDAR_INGEST_ENABLED=true."
        ),
    )

    feature(
        doc,
        "Plaud transcripts",
        "Plaud recordings flow into Supabase via webhook → API pull. Transcripts are "
        "linked to the right project (by attendee match against Harvest assignments), "
        "auto-embedded into RAG, and folded into that project's Brain as Recent "
        "decisions / Watchlist items.",
        requires="PLAUD_WEBHOOK_SECRET, PLAUD_API_KEY, PLAUD_INGEST_ENABLED=true.",
        notes="Activation gated on Plaud dev access (their gated form).",
    )

    # ── 4. Pipelines ───────────────────────────────────────
    add_h1(doc, "4. Pipelines")

    feature(
        doc,
        "Delivery pipeline (distributed transcoding)",
        "Drop a file into /Delivery-Queue/ on Dropbox. Kit posts a profile-selection "
        "prompt to the delivery channel. A render worker (FFmpeg on a studio PC) claims "
        "the job and transcodes to broadcast specs (ProRes, loudness normalization, "
        "channel mapping, naming convention). Output appears in /Delivery-Queue/.../delivery/. "
        "Multi-worker fleet with primary + fallback priorities and stale-job reassignment.",
        trigger="Drop file in Dropbox, or /kit deliver <path>. Status via /kit deliver status.",
        requires="DELIVERY_NOTIFY_CHANNEL_ID, FFmpeg + kit-render-worker installed on studio PCs.",
    )

    feature(
        doc,
        "Accessibility pipeline (captions + descriptive audio)",
        "Drop a video into /Accessibility-Queue/ on Dropbox. Kit produces .srt / .ttml / "
        ".txt caption tracks (Whisper) and optionally a descriptive-audio .mp3 narration "
        "(GPT-4 Vision + ElevenLabs TTS, sized to fit natural pauses in the video). "
        "Outputs land in a subfolder named after the video.",
        trigger="Drop video in Dropbox.",
        requires=(
            "OPENAI_API_KEY (captions + vision). Optional ELEVENLABS_API_KEY + "
            "ELEVENLABS_VOICE_ID + ACCESSIBILITY_NOTIFY_CHANNEL_ID for descriptive audio."
        ),
    )

    # ── 5. Studio Knowledge ────────────────────────────────
    add_h1(doc, "5. Studio Knowledge (RAG over project history)")

    feature(
        doc,
        "Project history Q&A",
        "Ask Kit about any project the studio has ever done: brief, budget, status, "
        "people involved, deliverables, what worked, what didn't. Backed by OpenAI "
        "text-embedding-3-small + pgvector + Supabase match_documents RPC.",
        trigger="@Kit who's the producer on STUDIO100? · @Kit biggest Microsoft project this year? · @Kit what was the brief for the Nike sizzle?",
        requires="OPENAI_API_KEY, plus a one-time Harvest project + client backfill.",
    )

    feature(
        doc,
        "Contacts lookup",
        "Search every client's contact list by name, email, or title. Returns the match "
        "with their client and role.",
        trigger="@Kit who do we talk to at Microsoft? · @Kit find Sarah Chen",
    )

    feature(
        doc,
        "Recent projects / clients",
        "Lists the most-recent projects (by start date) or top clients (by lifetime revenue).",
        trigger="@Kit what have we worked on lately? · @Kit who are our biggest clients?",
    )

    feature(
        doc,
        "Nightly auto-summarization (optional)",
        "Each project's summary doc gets regenerated nightly from the latest notes, "
        "transcripts, and decisions via Claude Haiku. Keeps narrative summaries current.",
        requires="STUDIO_KNOWLEDGE_AUTO_SUMMARIZE_ENABLED=true. Recommend waiting 2–3 weeks "
        "of accumulated activity before enabling.",
    )

    # ── 6. Brain ───────────────────────────────────────────
    add_h1(doc, "6. Brain (per-channel project intelligence)")

    add_para(
        doc,
        "Each project channel gets a living, versioned markdown brain mirrored to a "
        "Slack Canvas. The brain captures Operating context, Conventions & specs, "
        "Open decisions, Recent decisions, Watchlist, People & roles, and Glossary. "
        "Every bullet carries an inline source tag so any fact is traceable.",
    )

    feature(
        doc,
        "Brain canvas",
        "Open the Brain tab in any project channel to see the current state. Updates in "
        "place — there is one canvas per brain, edited as patches land. Provenance tags "
        "are stripped for readability but stored in the markdown. NEW brains default to "
        "producers_only (no canvas) for privacy; producers flip to team-visible with "
        "/kit brain visibility team.",
        trigger="/kit brain (producer/admin) · auto-seeded on /kit newproject completion (producers_only by default).",
    )

    feature(
        doc,
        "Brain visibility toggle",
        "Per-brain privacy flag. 'producers_only' (default) means no channel canvas — "
        "the brain is accessible only via /kit brain ephemeral text to producer+ users. "
        "'team' means a Slack Canvas tab is created and refreshed in the project channel, "
        "visible to anyone in the channel.",
        trigger="/kit brain visibility team | /kit brain visibility producers_only",
        requires="Producer or admin tier.",
    )

    feature(
        doc,
        "Self-learning writer",
        "Every substantive channel message and every saved note flows through a cheap "
        "deterministic classifier and then Claude Haiku, which proposes structured "
        "patches against the brain's section anchors. High-confidence patches (≥ 0.7) "
        "auto-apply. The canvas refreshes within seconds.",
        notes="Always on. Cost per surviving message ≈ a fraction of a cent.",
    )

    feature(
        doc,
        "Sourced in-thread answers",
        "When you @-mention Kit in a brain-bearing channel, the answer prefers the "
        "brain's own sections over generic project docs and ends with a visible "
        "_Sources:_ line built from the provenance tags of the bullets it used.",
    )

    feature(
        doc,
        "Why-lookup",
        "Ask which source backs any claim in the brain.",
        trigger="/kit brain why <claim>",
    )

    feature(
        doc,
        "Mistake-catch (narrow)",
        "On every message, Kit checks for contradictions against the brain's Glossary "
        "and the delivery date / spec lines in Operating context. High-confidence "
        "(≥ 0.85) catches reply in-thread with the correction and source. Questions and "
        "decisions are explicitly NOT flagged.",
        notes="Always on. Scope is intentionally narrow until trust is earned.",
    )

    feature(
        doc,
        "Deadline watch (cron)",
        "Hourly sweep walks every brain's Watchlist; items due within the lead window "
        "(default 3 days) get a ⚠️ in-channel flag; past-due items get 🚨. Dedup so "
        "each item flags once per window.",
        requires="KIT_BRAIN_DEADLINE_SWEEP_ENABLED=true. Lead window via KIT_BRAIN_DEADLINE_LEAD_DAYS.",
    )

    feature(
        doc,
        "Scavenger + DM approval (gated)",
        "Daily, Kit searches across the studio's other project docs for items related "
        "to each brain's Open decisions and Watchlist. The channel creator gets a DM "
        "with approve / reject buttons for each candidate. On approval, the item is "
        "folded into the brain with source preserved. Cross-channel context donation "
        "always asks, even under autonomous mode.",
        requires="KIT_BRAIN_SCAVENGER_ENABLED=true.",
    )

    feature(
        doc,
        "Nightly consolidator (gated)",
        "Keeps the brain tight: ages out stale Watchlist items past their grace window, "
        "compresses the Recent decisions log into a separate Earlier-decisions section "
        "after N entries, runs a conservative Haiku dedup pass on bullet-heavy sections.",
        requires="KIT_BRAIN_CONSOLIDATOR_ENABLED=true. Recommend a few weeks of accumulated "
        "material before enabling.",
    )

    # ── 7. Infrastructure ──────────────────────────────────
    add_h1(doc, "7. Infrastructure")

    add_h3(doc, "Agent registry")
    add_para(
        doc,
        "Eight domain-expert agents (Harvest, Dropbox, Frame.io, Slack, Boords, Delivery, "
        "Studio Knowledge, Brain). Each declares its capabilities, required env vars, "
        "and whether each action mutates state. The orchestrator picks the right "
        "specialist for any natural-language request.",
    )

    add_h3(doc, "Access control (three-tier)")
    add_para(
        doc,
        "admin / producer / artist. Gateway rules block whole agent actions; field-level "
        "filtering strips sensitive fields from anything that slips through. Every "
        "specialist dispatch passes through enforceAccess. Unknown users fail-closed to "
        "artist tier (never bypass).",
    )
    add_para(doc, "Tier matrix (high level):", bold=True)
    add_bullet(doc, "admin — sees everything (founder/owner)")
    add_bullet(
        doc,
        "producer — sees budgets/clients/dates/brief/contacts (per-project financial flag "
        "for raw budget queries); no margins/rates",
    )
    add_bullet(
        doc,
        "artist — name + project_code + harvest_project_id + status only. No Brain, no "
        "Studio Knowledge, no client/contact/budget/date data. Can log own time, see "
        "project tasks.",
    )

    add_h3(doc, "Role assignment")
    add_para(
        doc,
        "Admins assign roles via /kit role @user producer|artist|admin|freelancer. "
        "Stored in team_members.role; takes effect immediately. Omit the role to "
        "query someone's current setting.",
    )

    add_h3(doc, "Autonomy gates (always-on hard rules)")
    add_para(
        doc,
        "Two rules apply regardless of autonomy mode:",
    )
    add_bullet(doc, "Client-facing sends (emails, SOWs, scope alerts) are draft-only — never auto-sent.")
    add_bullet(doc, "Cross-channel context donation (scavenger) always asks for human approval.")

    add_h3(doc, "Scheduled jobs (Inngest + node-cron)")
    add_bullet(doc, "Pre-meeting briefing scan — every 15 min")
    add_bullet(doc, "Delivery Dropbox watcher — every 60 s")
    add_bullet(doc, "Accessibility Dropbox watcher — every 60 s")
    add_bullet(doc, "Studio Knowledge nightly summary — 9 am UTC (gated)")
    add_bullet(doc, "Harvest weekly resync — Sunday")
    add_bullet(doc, "Brain deadline sweep — hourly (gated)")
    add_bullet(doc, "Brain scavenger scan — daily 7 am UTC (gated)")
    add_bullet(doc, "Brain scavenger DM dispatch — daily 7:15 am UTC (gated)")
    add_bullet(doc, "Brain consolidator — nightly 10 am UTC (gated)")
    add_bullet(doc, "Daily hours check-in — 5 pm local Mon–Fri")
    add_bullet(doc, "Hours check-in nudge — 10 pm local Mon–Fri")

    # ── 8. Slash commands ──────────────────────────────────
    add_h1(doc, "8. Slash command reference")
    cmds = [
        ("/kit newproject", "Open the new-project intake card"),
        ("/kit onboard", "Onboard a freelancer (PM / CD / admin only)"),
        ("/kit status <name>", "Quick project lookup via Harvest"),
        ("/kit shotlist <script>", "Build a shot list canvas in the current channel"),
        ("/kit note [project | body]", "Save a freeform note to a project (or current channel's project)"),
        ("/kit deliver [path]", "Submit a transcode job. /kit deliver status for the queue."),
        ("/kit profiles · /kit profiles create", "List or create delivery profiles"),
        ("/kit workers · /kit workers opt-out <host>", "Render worker fleet management"),
        ("/kit access status", "Recent accessibility jobs"),
        ("/kit brain", "Open or refresh the current channel's Brain (producer/admin only)"),
        ("/kit brain why <claim>", "Show the sources behind a fact in the Brain"),
        ("/kit brain visibility team|producers_only", "Producer toggle for channel-canvas vs ephemeral-text mode"),
        ("/kit role @user producer|artist|admin|freelancer", "Admin only: assign a role. Omit the role to query."),
        ("/storyboard", "Storyboard intake (script paste or file drop)"),
        ("/kit help", "List all commands"),
    ]
    for cmd, desc in cmds:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(cmd)
        r1.bold = True
        r1.font.name = "Consolas"
        r1.font.size = Pt(10)
        r2 = p.add_run(f"  — {desc}")
        r2.font.size = Pt(11)

    # ── 9. Operator notes ──────────────────────────────────
    add_h1(doc, "9. Operator notes")
    add_para(
        doc,
        "See OPERATOR-TODO.md in the repo root for the full activation checklist "
        "(API keys, backfills, render worker installs, optional flag flips).",
    )
    add_para(doc, "Active deployment surfaces:", bold=True)
    add_bullet(doc, "Bolt app on Railway (Socket Mode, always-on)")
    add_bullet(doc, "Next.js + Inngest on Vercel (background jobs)")
    add_bullet(doc, "Supabase Postgres + pgvector (state + RAG)")
    add_bullet(doc, "Render workers on studio PCs (FFmpeg, primary + fallbacks)")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
